"""
Convert a Markdown document in docs/ into a formatted Word document.

Written because pandoc is not available in this environment. It handles the
subset of Markdown actually used by the proposal and report: headings,
paragraphs, bullet and numbered lists, tables, images, bold/italic inline
formatting, and horizontal rules.

Output is formatted to the conventions expected of an academic submission:
Times New Roman 12pt, 1.5 line spacing, justified body text, numbered figure
captions and a hanging-indent reference list.

Run:  python src/build_docx.py docs/PROPOSAL.md docs/Project_Proposal.docx
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)

PROJECT_ROOT = Path(__file__).resolve().parent.parent


# ---------------------------------------------------------------------------
# Styling
# ---------------------------------------------------------------------------

def setup_styles(doc):
    """Apply document-wide typography suitable for an academic submission."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    for level, size, bold in [("Heading 1", 15, True),
                              ("Heading 2", 13, True),
                              ("Heading 3", 12, True)]:
        st = doc.styles[level]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        st.paragraph_format.space_before = Pt(14)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.2


def add_inline(paragraph, text):
    """
    Render inline Markdown (**bold**, *italic*, `code`) into runs.

    Splitting on a single regex keeps the ordering of mixed formatting intact,
    which naive sequential replacement does not.
    """
    for part in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", text):
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(10.5)
        else:
            run.text = part


# ---------------------------------------------------------------------------
# Block handlers
# ---------------------------------------------------------------------------

def add_table(doc, rows):
    """Render a Markdown pipe table."""
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    body = []
    for line in rows[2:]:                       # rows[1] is the --- separator
        if line.strip():
            body.append([c.strip() for c in line.strip("|").split("|")])

    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = True

    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", text))
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = BODY_FONT

    for row_cells in body:
        cells = table.add_row().cells
        for i, text in enumerate(row_cells[:len(header)]):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, text)
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.name = BODY_FONT

    doc.add_paragraph()


def add_image(doc, md_path, alt):
    """Insert an image with a centred caption taken from the alt text."""
    path = (PROJECT_ROOT / "docs" / md_path).resolve()
    if not path.exists():
        path = (PROJECT_ROOT / md_path.replace("../", "")).resolve()
    if not path.exists():
        print(f"  ! image not found: {md_path}")
        return

    doc.add_picture(str(path), width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    run = cap.add_run(alt)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.name = BODY_FONT


def gather_item(lines, i, marker):
    """
    Collect a list item together with any wrapped continuation lines.

    Markdown wraps long list items across several source lines. Without this,
    each wrapped line becomes its own paragraph and loses the item's indent,
    which is visible in the rendered document as text falling back to the
    left margin.
    """
    text = re.sub(marker, "", lines[i].strip())
    i += 1
    while i < len(lines):
        nxt = lines[i].strip()
        if not nxt or re.match(r"^(#{1,4}\s|[-*]\s|\d+\.\s|\||!\[|-{3,}$)", nxt):
            break
        text += " " + nxt
        i += 1
    return text, i


def add_reference(doc, text):
    """Reference entry with a hanging indent, as APA requires."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    pf.space_after = Pt(10)
    pf.line_spacing = 1.5
    add_inline(p, text)


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def convert(md_file, out_file):
    lines = Path(md_file).read_text(encoding="utf-8").split("\n")

    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)

    in_references = False
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank
        if not stripped:
            i += 1
            continue

        # Explicit page break marker.
        # Used to start front-matter sections and chapters on a fresh page, as
        # academic submission convention requires.
        if stripped == "<!-- pagebreak -->":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue

        # Horizontal rule -> spacing only (avoids stray table borders)
        if re.fullmatch(r"-{3,}", stripped):
            i += 1
            continue

        # Image
        m = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", stripped)
        if m:
            add_image(doc, m.group(2), m.group(1))
            i += 1
            continue

        # Table
        if stripped.startswith("|") and i + 1 < len(lines) \
                and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, block)
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level, text = len(m.group(1)), m.group(2)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            in_references = "References" in text

            if level == 1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(16)
                run.font.name = BODY_FONT
            else:
                h = doc.add_heading(text, level=min(level - 1, 3))
                for run in h.runs:
                    run.font.name = BODY_FONT
                    run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", stripped):
            text, i = gather_item(lines, i, r"^[-*]\s+")
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Inches(0.5)
            pf.first_line_indent = Inches(-0.25)
            pf.line_spacing = 1.5
            pf.space_after = Pt(4)
            add_inline(p, "•\t" + text)
            continue

        # Numbered list.
        # Numbering is written literally rather than using Word's "List Number"
        # style, because that style continues its count from the previous
        # numbered list anywhere in the document -- so §9.4 would start at 6
        # instead of 1. Literal numbers keep each list independent.
        if re.match(r"^\d+\.\s+", stripped):
            num = re.match(r"^(\d+)\.", stripped).group(1)
            text, i = gather_item(lines, i, r"^\d+\.\s+")
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Inches(0.5)
            pf.first_line_indent = Inches(-0.35)
            pf.line_spacing = 1.5
            pf.space_after = Pt(4)
            add_inline(p, f"{num}.\t" + text)
            continue

        # Paragraph: gather continuation lines
        block = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and \
                not re.match(r"^(#{1,4}\s|[-*]\s|\d+\.\s|\||!\[|-{3,}$)",
                             lines[i].strip()):
            block.append(lines[i].strip())
            i += 1
        text = " ".join(block)

        if in_references:
            add_reference(doc, text)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, text)

    doc.save(out_file)
    print(f"Wrote {out_file}")


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "docs/PROPOSAL.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else "docs/Project_Proposal.docx"
    convert(PROJECT_ROOT / src, PROJECT_ROOT / dst)
